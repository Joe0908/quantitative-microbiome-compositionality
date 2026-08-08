"""Build polished DOCX deliverables from the audited manuscript sources.

Design preset: ``narrative_proposal``.
Header patterns: a restrained ``editorial_cover`` variant for the manuscript
and ``memo_masthead`` for the research-audit package.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
MANUSCRIPT_MD = ROOT / "manuscript.md"
AUDIT_MD = ROOT / "research_audit_and_submission.md"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "6B7280"
LIGHT = "F4F6F9"
GRID = "C9D1D9"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_dxa)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in [("top", 80), ("bottom", 80), ("start", 120), ("end", 120)]:
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width)))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


URL_RE = re.compile(r"https?://[^\s)]+")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*|https?://[^\s)]+)")


def add_inline(paragraph, text, size=11, color="000000"):
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("http"):
            trailing = ""
            while token and token[-1] in ".,;:":
                trailing = token[-1] + trailing
                token = token[:-1]
            add_hyperlink(paragraph, token, token)
            if trailing:
                run = paragraph.add_run(trailing)
                set_run_font(run, size=size, color=color)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=size - 0.5, color=DARK_BLUE)
            run.font.highlight_color = None
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=color, italic=True)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=9, color=MUTED)


def add_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs, default=0) + 1
    next_num = max(existing_num, default=0) + 1
    ids = {}
    for kind, fmt, text_value in [("bullet", "bullet", "•"), ("decimal", "decimal", "%1.")]:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(next_abs))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        lvl.append(p_pr)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_num))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(next_abs))
        num.append(abstract_ref)
        numbering.append(num)
        ids[kind] = next_num
        next_abs += 1
        next_num += 1
    return ids


def apply_list_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208


def configure_document(doc, running_title, first_page_label):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    add_inline(hp, running_title, size=9, color=MUTED)
    p_pr = hp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "D7DBE2")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(fp)

    doc.core_properties.title = running_title
    doc.core_properties.subject = first_page_label
    doc.core_properties.comments = "Generated from audited quantitative-microbiome analysis outputs."


def title_page(doc, title, subtitle, label, status_note):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label.upper())
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, size=23, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, size=13, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(status_note)
    set_run_font(r, size=10.5, color=RED, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared from the public LCPM and MetaCardis analysis pipeline")
    set_run_font(r, size=10.5, color=MUTED, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Working draft • {date.today().strftime('%d %B %Y')}")
    set_run_font(r, size=10, color=MUTED)
    doc.add_page_break()


def parse_table(lines, index):
    rows = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(set(cell) <= set("-: ") for cell in rows[1]):
        rows.pop(1)
    return rows, index


def column_widths(rows):
    n = len(rows[0])
    weights = []
    for col in range(n):
        lengths = [len(row[col]) if col < len(row) else 0 for row in rows]
        weights.append(max(8, min(55, max(lengths))))
    total = sum(weights)
    widths = [max(700, round(9360 * weight / total)) for weight in weights]
    diff = 9360 - sum(widths)
    widths[-1] += diff
    if widths[-1] < 500:
        shortage = 500 - widths[-1]
        widths[-1] = 500
        widths[0] -= shortage
    return widths


def add_table(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    widths = column_widths(rows)
    set_table_geometry(table, widths)
    for row_index, row in enumerate(rows):
        tr_pr = table.rows[row_index]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if row_index == 0:
            repeat_header = OxmlElement("w:tblHeader")
            repeat_header.set(qn("w:val"), "true")
            tr_pr.append(repeat_header)
        for col_index in range(cols):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = row[col_index] if col_index < len(row) else ""
            add_inline(p, text, size=(8 if cols <= 4 else 7.2))
            if row_index == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_markdown(doc, path, skip_first_title=False):
    lines = path.read_text(encoding="utf-8").splitlines()
    numbering_ids = add_numbering(doc)
    index = 0
    paragraph_buffer = []
    in_equation = False
    equation_lines = []

    def flush_buffer():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = " ".join(part.strip() for part in paragraph_buffer).strip()
            if text:
                p = doc.add_paragraph()
                add_inline(p, text)
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if in_equation:
            if stripped == "\\]":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                r = p.add_run(" ".join(equation_lines))
                set_run_font(r, name="Cambria Math", size=11, color=NAVY, italic=True)
                in_equation = False
                equation_lines = []
            else:
                equation_lines.append(stripped)
            index += 1
            continue
        if stripped == "\\[":
            flush_buffer()
            in_equation = True
            index += 1
            continue
        if not stripped:
            flush_buffer()
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and lines[index + 1].strip().startswith("|"):
            flush_buffer()
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            flush_buffer()
            level = len(heading_match.group(1))
            text_value = heading_match.group(2).replace("**", "")
            if level == 1 and skip_first_title:
                skip_first_title = False
            else:
                style_level = min(level, 3)
                p = doc.add_paragraph(style=f"Heading {style_level}")
                if level == 4:
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(3)
                add_inline(
                    p,
                    text_value,
                    size={1: 16, 2: 13, 3: 12, 4: 11.5}[level],
                    color={1: BLUE, 2: BLUE, 3: DARK_BLUE, 4: DARK_BLUE}[level],
                )
                if level >= 2 or text_value == "References":
                    numbering_ids = add_numbering(doc)
            index += 1
            continue
        quote_match = re.match(r"^>\s+(.+)$", stripped)
        if quote_match:
            flush_buffer()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "12")
            left.set(qn("w:space"), "6")
            left.set(qn("w:color"), "2E74B5")
            p_bdr.append(left)
            p_pr.append(p_bdr)
            add_inline(p, quote_match.group(1), color=DARK_BLUE)
            for run in p.runs:
                run.italic = True
            index += 1
            continue
        list_match = re.match(r"^-\s+(.+)$", stripped)
        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        check_match = re.match(r"^-\s+\[([ xX])\]\s+(.+)$", stripped)
        if check_match:
            flush_buffer()
            p = doc.add_paragraph()
            marker = "☒" if check_match.group(1).lower() == "x" else "☐"
            add_inline(p, f"{marker} {check_match.group(2)}")
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.space_after = Pt(4)
            index += 1
            continue
        if list_match:
            flush_buffer()
            p = doc.add_paragraph()
            apply_list_numbering(p, numbering_ids["bullet"])
            add_inline(p, list_match.group(1))
            index += 1
            continue
        if number_match:
            flush_buffer()
            p = doc.add_paragraph()
            apply_list_numbering(p, numbering_ids["decimal"])
            add_inline(p, number_match.group(1))
            index += 1
            continue
        paragraph_buffer.append(stripped)
        index += 1
    flush_buffer()


def append_figures(doc):
    doc.add_page_break()
    h = doc.add_paragraph(style="Heading 1")
    add_inline(h, "Main Figures", size=16, color=BLUE)
    figures = [
        ("Figure_1_study_design.png", "Figure 1. Study design and paired abundance representations."),
        ("Figure_2_prediction.png", "Figure 2. Repeated-cross-validation discrimination across abundance representations."),
        ("Figure_3_effect_concordance.png", "Figure 3. Concordance of QMP and row-closed feature effects."),
        ("Figure_4_robustness.png", "Figure 4. Filtering, covariate, and clinical-model sensitivity."),
        ("Supplementary_Figure_1_shared_species.png", "Supplementary Figure 1. Cross-cohort representation sensitivity across exact shared species."),
        ("Supplementary_Figure_2_prevalence_sensitivity.png", "Supplementary Figure 2. LCPM prevalence-threshold and test sensitivity."),
        ("Supplementary_Figure_3_total_load.png", "Supplementary Figure 3. Total microbial-load sensitivity."),
    ]
    for filename, caption_text in figures:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(ROOT / "figures" / filename), width=Inches(6.25))
        c = doc.add_paragraph(style="Caption")
        add_inline(c, caption_text, size=9, color=MUTED)
        c.paragraph_format.keep_with_next = False
        if filename in {"Figure_1_study_design.png", "Figure_3_effect_concordance.png"}:
            doc.add_page_break()


def audit_docx_structure(doc):
    section = doc.sections[0]
    assert round(section.page_width.inches, 3) == 8.5
    assert round(section.page_height.inches, 3) == 11
    for margin in [section.top_margin, section.right_margin, section.bottom_margin, section.left_margin]:
        assert round(margin.inches, 3) == 1.0
    assert round(section.header_distance.inches, 3) == 0.492
    assert round(section.footer_distance.inches, 3) == 0.492
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        assert tbl_pr.find(qn("w:tblW")).get(qn("w:w")) == "9360"
        assert tbl_pr.find(qn("w:tblInd")).get(qn("w:w")) == "120"


def build_manuscript():
    doc = Document()
    configure_document(doc, "Quantitative versus compositional microbiome representations", "Manuscript draft")
    title_page(
        doc,
        "A paired benchmark of quantitative, row-closed, and log-ratio gut microbiome profiles in two public cohorts",
        "A paired reanalysis of LCPM and MetaCardis",
        "Original Research Manuscript",
        "DRAFT FOR COAUTHOR AND STATISTICAL REVIEW — NOT YET SUBMISSION-READY",
    )
    add_markdown(doc, MANUSCRIPT_MD, skip_first_title=True)
    append_figures(doc)
    audit_docx_structure(doc)
    output = ROOT / "Quantitative_Microbiome_Manuscript_Draft.docx"
    doc.save(output)
    return output


def build_audit():
    doc = Document()
    configure_document(doc, "Research Audit and Submission Package", "Research audit")
    title_page(
        doc,
        "Research Audit and Submission Package",
        "Quantitative microbiome compositionality project",
        "Research Integrity • Reviewer Mode • Journal Strategy",
        "OBJECTIVE STATUS: PUBLISHABLE AFTER MODERATE REVISION",
    )
    add_markdown(doc, AUDIT_MD, skip_first_title=True)
    audit_docx_structure(doc)
    output = ROOT / "Research_Audit_and_Submission_Package.docx"
    doc.save(output)
    return output


if __name__ == "__main__":
    manuscript = build_manuscript()
    audit = build_audit()
    print(manuscript)
    print(audit)
